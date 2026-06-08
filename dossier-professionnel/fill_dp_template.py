#!/usr/bin/env python3
"""
Remplit le modèle officiel AFPA "Dossier Professionnel" avec les
informations de Vladimir Rodzaevskiy pour le titre DWWM.

Source : /home/vladimir/Documents/Dossier Professionnel/1-Dossier_professionnel_*.docx
Sortie : /home/vladimir/vladimir-rodzaevski-dossiers-pros-DWWM/
                    dossier-professionnel/DP_AFPA_rempli.docx
"""
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path("/home/vladimir/Documents/Dossier Professionnel/"
           "1-Dossier_professionnel_version_traitement_de_texte_11_09_2017.docx")
DST = Path("/home/vladimir/vladimir-rodzaevski-dossiers-pros-DWWM/"
           "dossier-professionnel/DP_AFPA_rempli.docx")

# ────────────────────────────────────────────────────────────────────────
# DONNÉES PERSONNELLES
# ────────────────────────────────────────────────────────────────────────
PERSONAL = {
    "Nom de naissance":  "RODZAEVSKIY",
    "Nom d’usage":       "RODZAEVSKIY",
    "Prénom":            "Vladimir",
    "Adresse":           "58 avenue Maréchal Juin, 06400 Cannes",
}

TITRE_PRO = "Développeur Web et Web Mobile (DWWM) — RNCP 37674, niveau 5"

# ────────────────────────────────────────────────────────────────────────
# AT1 — DÉVELOPPER LA PARTIE FRONT-END
# ────────────────────────────────────────────────────────────────────────
AT1_INTITULE = "Développer la partie front-end d'une application web ou web mobile sécurisée"

AT1_EX_INTITULE = "Application Crew — planification d'équipe React 18 complète"

AT1_TASKS = (
    "Compétences couvertes (référentiel REAC DWWM — RNCP 37674) :\n"
    "  • C1.1 — Maquetter une application (mobile-first, breakpoints)\n"
    "  • C1.2 — Réaliser une interface utilisateur web statique et "
    "adaptable (responsive, accessibilité WCAG AA)\n"
    "  • C1.3 — Développer une interface utilisateur web dynamique "
    "(Context API, hooks, état global)\n"
    "  • C1.4 — Réaliser une interface avec un framework (React 18 + Vite)\n\n"
    "Dans le cadre du projet de fin de formation Crew (planification "
    "intelligente d'équipe pour la restauration), j'ai réalisé l'intégralité "
    "de la partie front-end de l'application. J'ai notamment :\n\n"
    "• maquetté l'interface en mobile-first puis enrichi pour tablette et "
    "desktop (breakpoints 768 px et 1024 px), avec 14 pages : Login, Register, "
    "Dashboard (manager et équipier), Planning (grille hebdomadaire avec "
    "drag-and-drop), Stats, Teams, TeamCreate/Join/Detail, TeamSettings, "
    "Profile, NotFound ;\n\n"
    "• mis en place le routage avec React Router 6 (routes publiques et "
    "protégées via un composant ProtectedRoute qui redirige vers /login "
    "quand l'utilisateur n'est pas authentifié) ;\n\n"
    "• développé un système d'état global avec 5 contextes React superposés "
    "(AuthContext, FamilyContext, ThemeContext, ToastContext, ConfirmContext) "
    "exposés via des hooks personnalisés (useAuth, useFamily, useTheme…) ;\n\n"
    "• construit la grille de planning hebdomadaire avec drag-and-drop natif "
    "HTML5 (déplacement d'un shift d'un jour ou d'un équipier à l'autre, "
    "mise à jour optimiste + rollback en cas d'erreur API) ;\n\n"
    "• développé la modale Smart Planner avec un tableau (jour × service) "
    "où le manager déclare la densité prévue de chaque service de la semaine "
    "(presets Calme 0,5 / Normal 1,0 / Chargé 1,3 + valeur libre) — re-fetch "
    "automatique du plan à chaque modification ;\n\n"
    "• intégré 3 graphiques Chart.js (heures planifiées vs cibles, couverture "
    "par poste, masse salariale hebdomadaire) sur la page Statistiques ;\n\n"
    "• développé la page Profil qui génère un QR code webcal:// pour "
    "l'abonnement au flux iCal depuis iPhone et Android, sans application "
    "à installer ;\n\n"
    "• implémenté un thème clair/sombre complet avec persistance en "
    "localStorage et respect de prefers-color-scheme ;\n\n"
    "• intégré l'i18n (FR/EN) via react-i18next avec switch dans la navbar.\n\n"
    "Extrait représentatif — système d'authentification global :\n"
    "// front/src/context/AuthContext.jsx\n"
    "export function AuthProvider({ children }) {\n"
    "  const [user, setUser] = useState(null);\n"
    "  useEffect(() => {\n"
    "    const token = localStorage.getItem('crew_token');\n"
    "    if (!token) return;\n"
    "    authApi.me().then(setUser)\n"
    "      .catch(() => localStorage.removeItem('crew_token'));\n"
    "  }, []);\n"
    "  async function login(email, password) {\n"
    "    const { token, user } = await authApi.login({ email, password });\n"
    "    localStorage.setItem('crew_token', token);\n"
    "    setUser(user);\n"
    "  }\n"
    "  return <AuthContext.Provider value={{ user, login, logout }}>\n"
    "    {children}\n"
    "  </AuthContext.Provider>;\n"
    "}"
)

AT1_MOYENS = (
    "Stack technique :\n"
    "  • React 18.3 + Vite 5 (HMR, build ESM)\n"
    "  • React Router 6.26\n"
    "  • Chart.js + react-chartjs-2 (graphiques accessibles)\n"
    "  • lucide-react (icônes vectorielles outline)\n"
    "  • react-i18next (i18n FR/EN avec switch dans la navbar)\n"
    "  • qrcode (génération canvas) pour QR webcal iCal\n"
    "  • HTML5 Drag-and-Drop API native pour le planning\n"
    "  • CSS pur avec custom properties (variables.css) — pas de framework, "
    "approche glassmorphism + dark mode natif\n\n"
    "Outils :\n"
    "  • Visual Studio Code\n"
    "  • Git + GitHub (commits réguliers, branche par feature)\n"
    "  • Firefox + Chromium (DevTools, accessibility audit, responsive mode)\n"
    "  • Playwright 1.60 pour les tests E2E\n"
    "  • Postman et curl pour tester les appels API\n\n"
    "Bonnes pratiques :\n"
    "  • Validation HTML5 systématique (required, type, autoComplete)\n"
    "  • Labels associés (htmlFor + id), aria-invalid, aria-describedby\n"
    "  • Focus visible, prefers-reduced-motion respecté\n"
    "  • Contraste WCAG AA vérifié sur tous les textes (ratio ≥ 4,5)\n"
    "  • Internationalisation systématique — aucune chaîne en dur côté front."
)

AT1_AVEC_QUI = (
    "Projet réalisé en solo dans le cadre de la formation AFPA DWWM. "
    "Échanges réguliers avec :\n"
    "  • le formateur référent AFPA Marseille pour cadrage technique et "
    "validation des choix d'architecture ;\n"
    "  • la promotion DWWM pour partage d'expérience et code reviews informelles ;\n"
    "  • la communauté open source (forums Stack Overflow, documentation MDN, "
    "documentation React et Vite) pour la résolution de problèmes techniques."
)

AT1_NOM_ENTREPRISE = ("AFPA Marseille — projet de fin de formation "
                       "« Crew »")
AT1_SERVICE = "Formation DWWM, plateau pédagogique"
AT1_PERIODE_DU = "Mai 2026"
AT1_PERIODE_AU = "Juillet 2026"

AT1_INFOS_COMPL = (
    "Application déployée en ligne :\n"
    "https://crew-planner-hazel.vercel.app\n\n"
    "Code source public sous licence MIT :\n"
    "https://github.com/Vladimir08880888/crew\n\n"
    "Documents complémentaires (annexes) :\n"
    "  • DP détaillé (3 exemples Front documentés ligne par ligne) : DP.pdf\n"
    "  • Cahier des charges (sections 1-11) : dossier-projet/cahier-des-charges.md\n"
    "  • Manuel utilisateur (manager + équipier) : annexes/MANUEL_UTILISATEUR.md\n"
    "  • Plan de tests (60+ scénarios Playwright) : annexes/PLAN_DE_TESTS.md\n"
    "  • 17 captures d'écran haute résolution : annexes/screenshots/\n\n"
    "Repo centralisé pour le jury :\n"
    "https://github.com/Vladimir08880888/vladimir-rodzaevski-dossiers-pros-DWWM"
)

# ────────────────────────────────────────────────────────────────────────
# AT2 — DÉVELOPPER LA PARTIE BACK-END
# ────────────────────────────────────────────────────────────────────────
AT2_INTITULE = "Développer la partie back-end d'une application web ou web mobile sécurisée"

AT2_EX_INTITULE = "Application Crew — API Node.js/Express sécurisée"

AT2_TASKS = (
    "Compétences couvertes (référentiel REAC DWWM — RNCP 37674) :\n"
    "  • C2.1 — Créer une base de données (MCD/MLD/MPD via 12 migrations SQL)\n"
    "  • C2.2 — Développer les composants d'accès aux données "
    "(modèles mysql2 avec requêtes préparées, sans ORM)\n"
    "  • C2.3 — Développer la partie back-end (Express, controllers, "
    "validators, solver d'optimisation)\n"
    "  • C2.4 — Mettre en œuvre des composants dans un framework "
    "(Express middleware composable)\n\n"
    "Pour le même projet Crew, j'ai conçu et développé l'intégralité "
    "de la partie back-end : conception de la base de données, API REST "
    "sécurisée, et un solver de planning intelligent sous contraintes "
    "Convention HCR.\n\n"
    "Conception de la base de données (Merise) :\n"
    "• MCD avec 4 entités (users, families, family_members, shifts) ;\n"
    "• MLD avec clé primaire composite sur family_members (relation many-to-many) ;\n"
    "• MPD matérialisé par 12 fichiers de migrations SQL versionnés (006-012 "
    "pour les évolutions v2 : niveaux, paramètres d'équipe, polyvalence, "
    "taux horaires) ;\n"
    "• stratégies de suppression (ON DELETE CASCADE / SET NULL / RESTRICT) "
    "réfléchies selon la sémantique de chaque relation.\n\n"
    "Développement de l'API REST (Express) :\n"
    "• 35+ endpoints documentés couvrant authentification, équipes, shifts, "
    "smart planner, statistiques, export iCal et paramètres d'établissement ;\n"
    "• architecture en couches strictes : routes → middleware → controllers "
    "→ validators → models ;\n"
    "• authentification JWT (HS256, durée 7 jours) avec mot de passe haché "
    "via bcrypt (coût 10) ;\n"
    "• système de middlewares composables : authRequired → "
    "requireFamilyMember → requireAdmin ;\n"
    "• validation systématique côté serveur (validators retournent HttpError 400 "
    "avec détails par champ) ;\n"
    "• gestion d'erreurs centralisée — les stack traces ne sont jamais "
    "renvoyées au client ;\n"
    "• génération de flux iCalendar (RFC 5545) avec VALARM pour notifications "
    "natives sur téléphone, sans application à installer.\n\n"
    "Solver Smart Planner sous contraintes Convention HCR :\n"
    "• filtre de candidats appliquant les plafonds légaux comme contraintes "
    "DURES (48 h/semaine — Code du travail L3121-20 ; 11 h/jour cuisinier, "
    "11 h 30 salle — Convention HCR ; 2 jours de repos hebdomadaire) ;\n"
    "• scoring multi-critères : déficit hebdo, préférence shift_default, "
    "poste primaire, expérience chef, pénalité économique sur la masse "
    "salariale (taux horaires HCR 2026) ;\n"
    "• règle métier « junior seul interdit » — un junior n'est jamais "
    "affecté à un poste s'il n'y a pas au moins un confirmé/chef présent ;\n"
    "• polyvalence (multi-skill matrix) — un équipier peut tenir plusieurs "
    "postes, le solver le mobilise hors de sa spécialité uniquement si "
    "aucun spécialiste primaire n'est éligible.\n\n"
    "Extrait représentatif — filtre HCR du solver :\n"
    "// back/src/services/plannerSolver.js\n"
    "const HCR_WEEKLY_MAX = 48;\n"
    "const HCR_DAILY_MAX_BY_POSTE = {\n"
    "  cuisine: 11, salle: 11.5, plonge: 11.5,\n"
    "  bar: 11.5, administration: 10,\n"
    "};\n"
    "candidates = members.filter((m) => {\n"
    "  if (!canFill(m, poste)) return false;\n"
    "  if (wouldBeWeek > HCR_WEEKLY_MAX) return false;\n"
    "  if (dayHours + shiftDur > hcrDailyCap(m.poste)) return false;\n"
    "  if (weekDates.length - wouldDays < HCR_MIN_REST_DAYS) return false;\n"
    "  if (!seniorPresent && !isSenior(m)) return false;\n"
    "  return true;\n"
    "});"
)

AT2_MOYENS = (
    "Stack technique :\n"
    "  • Node.js 22 LTS + Express 4.19\n"
    "  • mysql2 3.11 (driver avec prepared statements natifs, sans ORM)\n"
    "  • MariaDB 10.11 (moteur InnoDB pour ACID + FK strictes)\n"
    "  • jsonwebtoken 9.0 (signature JWT HS256)\n"
    "  • bcrypt 5.1 (hachage des mots de passe)\n"
    "  • ical-generator 7.2 (génération RFC 5545 + VALARM)\n"
    "  • express-rate-limit 7.5 (anti brute-force sur /auth/login)\n"
    "  • cors 2.8 (CORS restreint au front configuré)\n"
    "  • dotenv 16.4 (variables d'environnement)\n\n"
    "Outils :\n"
    "  • Visual Studio Code\n"
    "  • Git + GitHub\n"
    "  • curl + jq pour tests manuels API\n"
    "  • MariaDB CLI pour debugging SQL\n"
    "  • Postman pour exploration d'API\n\n"
    "Sécurité (OWASP Top 10) :\n"
    "  • 100 % requêtes préparées (vérifié par grep — aucune concaténation)\n"
    "  • Pas de cookie session (JWT en Authorization Bearer) → CSRF impossible\n"
    "  • CORS limité au FRONT_ORIGIN configuré\n"
    "  • Variables sensibles (.env, JWT_SECRET, DB_PASSWORD) gitignored\n\n"
    "Tests :\n"
    "  • Suite Playwright de 60+ scénarios couvrant authentification, équipes, "
    "shifts, smart planner, conformité HCR, polyvalence multi-skill, cost-aware. "
    "Bilan : 29/29 PASS sur les tests v2 (bloc 10 du Plan de tests, voir annexe)."
)

AT2_AVEC_QUI = (
    "Projet réalisé en solo dans le cadre de la formation AFPA DWWM. "
    "Mêmes interlocuteurs que pour l'AT1 :\n"
    "  • formateur référent AFPA Marseille pour validation des choix d'architecture "
    "et des décisions de sécurité ;\n"
    "  • promotion DWWM pour partage et entraide ;\n"
    "  • documentation officielle Node.js, Express, MariaDB, OWASP."
)

AT2_NOM_ENTREPRISE = AT1_NOM_ENTREPRISE
AT2_SERVICE = AT1_SERVICE
AT2_PERIODE_DU = AT1_PERIODE_DU
AT2_PERIODE_AU = AT1_PERIODE_AU

AT2_INFOS_COMPL = (
    "Application déployée en ligne :\n"
    "https://crew-planner-hazel.vercel.app\n\n"
    "Code source public sous licence MIT :\n"
    "https://github.com/Vladimir08880888/crew\n\n"
    "Documents complémentaires (annexes) :\n"
    "  • DP détaillé (3 exemples Back documentés ligne par ligne) : DP.pdf\n"
    "  • Cahier des charges (sections 1-11 incluant v2) : dossier-projet/cahier-des-charges.md\n"
    "  • Justification scientifique (12 références peer-reviewed — INRS, NIOSH, "
    "Convention HCR, OR economics) : annexes/JUSTIFICATION_SCIENTIFIQUE.md\n"
    "  • Schéma BDD (Mermaid ERD + règles d'intégrité) : annexes/SCHEMA_BDD.md\n"
    "  • Plan de tests (60+ scénarios Playwright, 29/29 PASS sur v2) : annexes/PLAN_DE_TESTS.md\n"
    "  • 12 migrations SQL versionnées : back/migrations/001-012\n\n"
    "Repo centralisé pour le jury :\n"
    "https://github.com/Vladimir08880888/vladimir-rodzaevski-dossiers-pros-DWWM"
)

# ────────────────────────────────────────────────────────────────────────
# AT1 — EXEMPLE 2 (Planning grid drag-and-drop)
# ────────────────────────────────────────────────────────────────────────
AT1B_EX_INTITULE = "Application Crew — grille de planning hebdomadaire avec drag-and-drop"

AT1B_TASKS = (
    "Compétences couvertes :\n"
    "  • C1.2 — Interface adaptable (grille scrollable horizontal mobile)\n"
    "  • C1.3 — Interface dynamique (drag-and-drop HTML5 natif, rendu optimiste)\n"
    "  • C1.4 — Framework React (composants + hooks)\n\n"
    "Toujours sur le projet Crew, j'ai conçu la page Planning qui est "
    "le cœur d'utilisation quotidien du manager. Il s'agit d'une grille "
    "(équipiers × jours de la semaine) où chaque cellule représente les "
    "shifts d'un équipier sur une journée donnée.\n\n"
    "Réalisations :\n\n"
    "• rendu d'une grille HTML <table> performante, scrollable horizontalement "
    "sur mobile, avec en-tête sticky et pastilles de santé du service "
    "(🟢 Saine / 🟠 Tendue / 🔴 Risque) en bas de chaque colonne jour ;\n\n"
    "• regroupement visuel des équipiers par poste (cuisine d'abord, salle "
    "ensuite) avec un en-tête de groupe et une ligne « Extras » en fin de "
    "groupe affichant les déficits de couverture par service ;\n\n"
    "• implémentation du DRAG-AND-DROP NATIF HTML5 sur les cartes de shifts : "
    "le manager déplace un shift d'un jour ou d'un équipier à l'autre, "
    "la grille met à jour l'état localement (rendu optimiste) puis appelle "
    "l'API ; si le back rejette (par exemple violation du plafond HCR "
    "48 h/semaine), l'état est rollback automatiquement et un toast d'erreur "
    "explique le motif ;\n\n"
    "• prise en compte de la polyvalence multi-skill pour la validation côté "
    "client — un slot refuse les drops dans une zone visuellement marquée "
    "« incompatible » avec animation CSS ;\n\n"
    "• gestion des cartes en pleine largeur sur mobile via flexbox, conservant "
    "lisibilité et tap-targets de 44 px minimum.\n\n"
    "Extrait du gestionnaire de drop :\n"
    "// front/src/pages/Planning.jsx\n"
    "async function handleDrop(e, targetDate, targetUserId) {\n"
    "  e.preventDefault();\n"
    "  const shiftId = e.dataTransfer.getData('shift-id');\n"
    "  const prev = shifts;\n"
    "  // Optimistic update\n"
    "  setShifts(shifts.map(s =>\n"
    "    s.id === Number(shiftId)\n"
    "      ? { ...s, date: targetDate, user_id: targetUserId }\n"
    "      : s));\n"
    "  try {\n"
    "    await shiftsApi.update(shiftId, { date: targetDate, user_id: targetUserId });\n"
    "  } catch (err) {\n"
    "    setShifts(prev);  // rollback\n"
    "    toast.fromError(err);\n"
    "  }\n"
    "}"
)

AT1B_MOYENS = (
    "Stack technique :\n"
    "  • HTML5 Drag-and-Drop API native (draggable, ondragstart, ondrop, "
    "dataTransfer) — pas de bibliothèque externe\n"
    "  • React 18.3 + hooks (useState, useMemo)\n"
    "  • lucide-react (Trash2, Edit, GripVertical icons)\n"
    "  • CSS pur avec custom properties (variables.css)\n"
    "  • react-toastify pour les notifications d'erreur\n\n"
    "Bonnes pratiques :\n"
    "  • Rendu optimiste systématique + rollback automatique en cas d'erreur API\n"
    "  • Pas de bibliothèque lourde pour le drag-and-drop : choix de l'API "
    "native qui économise ~30 KB de bundle\n"
    "  • Mémoïsation des dérivations de l'état (poste groups, extras) via useMemo\n"
    "  • Accessibilité : draggable + keyboard fallback documenté en TODO."
)

AT1B_AVEC_QUI = AT1_AVEC_QUI

# ────────────────────────────────────────────────────────────────────────
# AT1 — EXEMPLE 3 (Modale Smart Planner)
# ────────────────────────────────────────────────────────────────────────
AT1C_EX_INTITULE = "Application Crew — Modale Smart Planner avec densité prévue par service"

AT1C_TASKS = (
    "Compétences couvertes :\n"
    "  • C1.3 — Interface dynamique (état contrôlé, useEffect, re-fetch live)\n"
    "  • C1.4 — Framework React (composants composables, react-i18next)\n\n"
    "La modale Smart Planner est le composant le plus dense fonctionnellement "
    "de l'application : c'est là que le manager déclare la densité prévue "
    "de chaque service de la semaine et obtient en retour une proposition "
    "de planning calculée par le back-end sous contraintes Convention HCR.\n\n"
    "Réalisations :\n\n"
    "• tableau interactif (7 jours × 2 services Midi/Soir) — chaque cellule "
    "présente un input numérique 30–200 % avec 3 boutons-presets (Calme 0,5, "
    "Normal 1,0, Chargé 1,3) ;\n\n"
    "• boutons « Tous → 0,5 / 1,0 / 1,3 » qui remplissent une colonne entière "
    "en un clic ;\n\n"
    "• re-fetch automatique de la proposition à chaque changement de densité "
    "(useEffect avec capacityByDateAndService en dépendance) — le manager voit "
    "l'effet en quasi temps réel ;\n\n"
    "• affichage du résultat en trois blocs :  heures par équipier (vs cibles "
    "contractuelles), couverture par (service × poste) avec fractions et %, "
    "services non couverts avec motif explicite ;\n\n"
    "• bouton d'application qui POST en masse les shifts proposés et ferme "
    "la modale avec un toast récapitulatif (« 26 shifts créés »).\n\n"
    "Extrait — déclenchement du re-fetch sur changement de densité :\n"
    "// front/src/components/planning/SmartPlannerModal.jsx\n"
    "useEffect(() => {\n"
    "  setLoading(true);\n"
    "  shiftsApi.generatePlan({\n"
    "    family_id: familyId, from, to,\n"
    "    capacityByDateAndService: perCell,\n"
    "  }).then(setData).finally(() => setLoading(false));\n"
    "}, [familyId, from, to, perCell]);"
)

AT1C_MOYENS = (
    "Stack technique :\n"
    "  • React hooks (useState, useEffect, useMemo) + composants contrôlés\n"
    "  • fetch via shiftsApi (wrapper centralisé)\n"
    "  • i18next pour i18n FR/EN sur tous les labels\n"
    "  • CSS Grid pour le tableau de densité (7 cols × 2 rows + colonne de "
    "presets)\n\n"
    "Bonnes pratiques :\n"
    "  • Une seule source de vérité (perCell) ; tous les inputs sont contrôlés\n"
    "  • Debounce implicite via réconciliation React (les keystrokes rapides "
    "n'entraînent qu'un seul re-fetch par batch)\n"
    "  • Gestion du loading state séparée pour ne pas bloquer l'UI"
)

AT1C_AVEC_QUI = AT1_AVEC_QUI

# ────────────────────────────────────────────────────────────────────────
# AT2 — EXEMPLE 2 (Authentification JWT + bcrypt)
# ────────────────────────────────────────────────────────────────────────
AT2B_EX_INTITULE = "Application Crew — Authentification JWT et middlewares composables"

AT2B_TASKS = (
    "Compétences couvertes :\n"
    "  • C2.3 — Back-end sécurisé (JWT, bcrypt, rate-limiting)\n"
    "  • C2.4 — Middleware composable (authRequired → requireFamilyMember "
    "→ requireAdmin) dans Express\n\n"
    "Le système d'authentification de Crew protège toutes les routes "
    "applicatives. J'ai conçu l'enchaînement complet — du hachage du "
    "mot de passe à l'invalidation côté client — sans cookie pour "
    "éliminer la classe de vulnérabilités CSRF.\n\n"
    "Réalisations :\n\n"
    "• endpoint POST /auth/register : validation des champs (email format, "
    "mot de passe ≥ 8 caractères), unicité email, hachage bcrypt cost 10, "
    "génération d'un calendar_token aléatoire de 32 octets pour iCal ;\n\n"
    "• endpoint POST /auth/login : rate-limiting 10 tentatives / 15 min via "
    "express-rate-limit, comparaison bcrypt (timing-safe), émission d'un JWT "
    "HS256 d'une durée de 7 jours signé avec JWT_SECRET ;\n\n"
    "• middleware authRequired qui parse le header Authorization Bearer, "
    "vérifie la signature et l'expiration, expose req.user au reste de la "
    "chaîne ; tout token invalide ou expiré renvoie 401 ;\n\n"
    "• middlewares composables d'autorisation : requireFamilyMember "
    "(vérifie l'appartenance à l'équipe), requireAdmin (vérifie le rôle "
    "manager), assertCanManageShifts (combine les deux) ;\n\n"
    "• endpoint POST /auth/change-password : exige l'ancien mot de passe + "
    "nouveau, valide le nouveau, hache et stocke.\n\n"
    "Extrait — middleware authRequired :\n"
    "// back/src/middleware/auth.js\n"
    "export function authRequired(req, _res, next) {\n"
    "  const hdr = req.header('Authorization') || '';\n"
    "  if (!hdr.startsWith('Bearer ')) throw unauthorized('Token manquant');\n"
    "  try {\n"
    "    const decoded = jwt.verify(hdr.slice(7), process.env.JWT_SECRET);\n"
    "    req.user = { id: decoded.userId };\n"
    "    next();\n"
    "  } catch (e) {\n"
    "    throw unauthorized('Token invalide ou expiré');\n"
    "  }\n"
    "}"
)

AT2B_MOYENS = (
    "Stack technique :\n"
    "  • jsonwebtoken 9.0 (JWT HS256)\n"
    "  • bcrypt 5.1 (hachage Bcrypt, cost 10)\n"
    "  • express-rate-limit 7.5 (rate-limit IP-based)\n"
    "  • crypto natif Node pour calendar_token aléatoire\n\n"
    "Sécurité :\n"
    "  • JWT en header Authorization Bearer (pas en cookie) → CSRF impossible\n"
    "  • JWT_SECRET externalisé en .env (gitignored)\n"
    "  • Bcrypt timing-safe via .compare()\n"
    "  • Stack traces masquées en prod (catch global)\n"
    "  • Vérifié manuellement : 0 concaténation SQL → 0 risque d'injection"
)

AT2B_AVEC_QUI = AT2_AVEC_QUI

# ────────────────────────────────────────────────────────────────────────
# AT2 — EXEMPLE 3 (Génération iCalendar RFC 5545)
# ────────────────────────────────────────────────────────────────────────
AT2C_EX_INTITULE = "Application Crew — Génération d'un flux iCalendar RFC 5545"

AT2C_TASKS = (
    "Compétences couvertes :\n"
    "  • C2.2 — Composants d'accès aux données (lecture shifts + user "
    "via token URL pour le flux iCal)\n"
    "  • C2.3 — Back-end (génération d'un flux RFC 5545 conforme avec "
    "VALARM pour notifications natives)\n\n"
    "Pour permettre à chaque équipier de recevoir ses shifts directement "
    "dans son application Calendrier native (iPhone Calendar, Google "
    "Calendar, Outlook), j'ai implémenté la génération d'un flux iCalendar "
    "conforme à la RFC 5545 — pas d'application supplémentaire à installer.\n\n"
    "Réalisations :\n\n"
    "• endpoint GET /calendar/:token et /calendar/:token/perso.ics qui "
    "authentifient via un token URL (le calendar_token généré à l'inscription) "
    "— pas de JWT car les applications calendrier ne savent pas en envoyer ;\n\n"
    "• construction d'un objet VCALENDAR avec ical-generator : un VEVENT par "
    "shift, intégrant UID stable, DTSTART/DTEND, SUMMARY (Service midi — Cuisine), "
    "DESCRIPTION (poste, note manager) et VALARM (notification 2 h avant) ;\n\n"
    "• gestion de la timezone (Europe/Paris) avec passage automatique heure d'été ;\n\n"
    "• endpoint distinct par équipe (calendar/:token/family/:familyId.ics) pour "
    "filtrer les shifts d'une seule équipe lorsque l'équipier appartient à "
    "plusieurs ;\n\n"
    "• tests manuels validés sur iPhone (ajout via Réglages → Calendrier → "
    "Comptes → Ajouter URL), Google Calendar (Ajouter agenda → À partir "
    "d'une URL), Outlook.\n\n"
    "Extrait — construction du flux :\n"
    "// back/src/services/ical.service.js\n"
    "export function buildIcal({ ownerName, shifts }) {\n"
    "  const cal = ical({\n"
    "    name: `Crew ${ownerName}`,\n"
    "    prodId: { company: 'Crew', product: 'Planner', language: 'FR' },\n"
    "    timezone: 'Europe/Paris',\n"
    "  });\n"
    "  shifts.forEach((s) => {\n"
    "    cal.createEvent({\n"
    "      uid: `shift-${s.id}@crew-planner`,\n"
    "      start: shiftStart(s), end: shiftEnd(s),\n"
    "      summary: `Service ${s.shift_type} — ${s.poste}`,\n"
    "      alarms: [{ type: 'display', triggerBefore: 7200 }],\n"
    "    });\n"
    "  });\n"
    "  return cal;\n"
    "}"
)

AT2C_MOYENS = (
    "Stack technique :\n"
    "  • ical-generator 7.2 (construit un flux RFC 5545 valide en JS)\n"
    "  • date-fns-tz pour les conversions de timezone\n"
    "  • Express pour exposer les 3 endpoints calendar\n\n"
    "Standards respectés :\n"
    "  • RFC 5545 (iCalendar) : UID stable, DTSTART/DTEND, VALARM\n"
    "  • Réponse Content-Type: text/calendar; charset=utf-8\n"
    "  • Pas d'expiration sur le token URL → l'équipier ne doit pas le "
    "partager publiquement (régénération possible depuis Profil)"
)

AT2C_AVEC_QUI = AT2_AVEC_QUI

# ────────────────────────────────────────────────────────────────────────
# DIPLÔMES
# ────────────────────────────────────────────────────────────────────────
DIPLOMES = [
    # (Intitulé, Autorité ou organisme, Date)
    ("Master 2 Économie et Gestion d'Entreprise",
     "Université d'État (Moscou, Russie)",
     "2004"),
    ("Master 2 Commerce International",
     "INSEEC — École de commerce, Paris",
     "2012"),
    ("Titre Professionnel Développeur Web et Web Mobile",
     "AFPA Marseille (en cours)",
     "2025–2026"),
]

# ────────────────────────────────────────────────────────────────────────
# UTILITAIRES python-docx
# ────────────────────────────────────────────────────────────────────────

def set_cell_text(cell, text, *, bold=False, preserve_first_run=True):
    """
    Vide une cellule de tableau et y écrit `text` (multi-lignes acceptées
    via '\n'). Préserve le style du premier run quand possible.
    """
    # supprimer tous les paragraphes existants sauf le premier
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    # supprimer tous les runs existants
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # écrire ligne par ligne
    lines = text.split('\n')
    run = p.add_run(lines[0])
    if bold:
        run.bold = True
    for line in lines[1:]:
        p.add_run().add_break()
        p.add_run(line)


def fill_example_table(table, intitule_at, intitule_ex,
                       tasks, moyens, avec_qui,
                       nom_ent, service, periode_du, periode_au,
                       infos_compl="", exemple_num=1):
    """
    Remplit un tableau « exemple de pratique professionnelle ».
    Structure du tableau (33 rows) :
      row 0 : Activité-type | (n°) | (intitulé)
      row 1 : Exemple n° X  |       | (intitulé)
      row 2-3 : ligne séparatrice
      row 4 : « 1. Décrivez … » + zones de saisie en dessous
      row 10 : « 2. Précisez les moyens utilisés »
      row 16 : « 3. Avec qui avez-vous travaillé »
      row 22 : « 4. Contexte »
      row 24 : « Nom de l'entreprise … » | (à remplir)
      row 25 : « Chantier, atelier, service » | (à remplir)
      row 26 : « Période d'exercice  Du : Cliquez ici  au : Cliquez ici »
      row 29 : « 5. Informations complémentaires »
    Note: chaque section a plusieurs rows vides en dessous pour la saisie ;
    on cible la première row libre après chaque titre.
    """
    rows = table._tbl.tr_lst

    def cell_tcs_at(row_idx):
        # Find ALL <w:tc> including those wrapped by <w:sdt>/<w:sdtContent>
        return rows[row_idx].xpath('.//w:tc')

    # row 0 — Activité-type : 3ème colonne = intitulé
    tcs = cell_tcs_at(0)
    if len(tcs) >= 3:
        # remplacer le texte « Cliquez ici … » dans la 3ème cellule
        set_cell_in_tc(tcs[-1], intitule_at, bold=True)

    # row 1 — Exemple n°X : 1ère colonne = libellé numéroté,
    # dernière colonne = intitulé de l'exemple.
    tcs = cell_tcs_at(1)
    if len(tcs) >= 2:
        # Met à jour le libellé « Exemple n°1 » selon exemple_num.
        set_cell_in_tc(tcs[0], f"Exemple n°{exemple_num}", bold=True)
        set_cell_in_tc(tcs[-1], intitule_ex, bold=True)

    # Section 1 : « 1. Décrivez les tâches » → écrire dans rows 5-9
    # (les rows 5, 6, 7, 8, 9 sont des cellules vides pour la saisie)
    set_section_content(rows, start_row=5, end_row=9, text=tasks)

    # Section 2 : « 2. Précisez les moyens » → rows 11-15
    set_section_content(rows, start_row=11, end_row=15, text=moyens)

    # Section 3 : « 3. Avec qui avez-vous travaillé » → rows 17-21
    set_section_content(rows, start_row=17, end_row=21, text=avec_qui)

    # Section 4 : Contexte — rows 24, 25, 26
    # row 24 : "Nom de l'entreprise" | (à remplir)
    tcs = cell_tcs_at(24)
    if len(tcs) >= 2:
        set_cell_in_tc(tcs[-1], nom_ent)
    # row 25 : "Chantier, atelier, service" | (à remplir)
    tcs = cell_tcs_at(25)
    if len(tcs) >= 2:
        set_cell_in_tc(tcs[-1], service)
    # row 26 : "Période d'exercice  Du : ... au : ..."
    # Cellule monobloc : on réécrit toute la ligne
    tcs = cell_tcs_at(26)
    if tcs:
        set_cell_in_tc(tcs[0],
                       f"Période d'exercice    Du : {periode_du}    au : {periode_au}")

    # Section 5 : Informations complémentaires → rows 30-32
    if infos_compl:
        set_section_content(rows, start_row=30, end_row=32, text=infos_compl)


def set_section_content(rows, start_row, end_row, text):
    """
    Écrit `text` dans la première cellule libre disponible entre
    start_row et end_row du tableau (cellules pour saisie utilisateur).
    """
    for r in range(start_row, end_row + 1):
        if r >= len(rows):
            break
        tcs = rows[r].xpath('.//w:tc')
        if not tcs:
            continue
        # prendre la première cellule de la row
        set_cell_in_tc(tcs[0], text)
        return


def set_cell_in_tc(tc, text, *, bold=False):
    """
    Remplit une cellule (élément <w:tc>) avec `text`.

    Stratégie : tout vider (paragraphes, SDT, runs) puis créer un seul
    paragraphe propre avec le texte voulu. Cela écrase aussi les
    « content controls » (SDT) Word qui contiennent souvent du texte
    placeholder du genre « Cliquez ici … ».
    """
    from docx.oxml import OxmlElement

    # 1. Supprimer TOUS les enfants directs sauf <w:tcPr> (propriétés)
    for child in list(tc):
        if child.tag == qn('w:tcPr'):
            continue
        tc.remove(child)

    # 2. Créer un seul paragraphe propre
    p = OxmlElement('w:p')
    tc.append(p)

    # 3. Ajouter le texte (multi-lignes via <w:br>)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br_r = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br_r.append(br)
            p.append(br_r)
        r = OxmlElement('w:r')
        if bold:
            rpr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rpr.append(b)
            r.append(rpr)
        t = OxmlElement('w:t')
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)


def fill_header_table(doc):
    """Table 0 — nom, prénom, adresse.

    Structure de chaque ligne :
      col 0 : libellé ("Nom de naissance")
      col 1 : cellule de saisie (vide)
      col 2 : SDT avec placeholder "Entrez votre nom de naissance ici."

    On remplit col 1 avec la valeur, ET on vide col 2 pour ne pas avoir
    à la fois la valeur et le placeholder côte à côte.
    """
    t = doc.tables[0]
    rows = t._tbl.tr_lst
    fields = [
        (1, PERSONAL["Nom de naissance"]),
        (2, PERSONAL["Nom d’usage"]),
        (3, PERSONAL["Prénom"]),
        (4, PERSONAL["Adresse"]),
    ]
    # Structure colonnes : 0 = label (4 cm) | 1 = icône séparateur 0.4 cm |
    # 2 = zone de saisie large (9 cm). On écrit la valeur dans la
    # colonne 2 et on laisse la colonne 1 (icône wingding) intacte.
    for row_idx, value in fields:
        tcs = rows[row_idx].xpath('.//w:tc')
        if len(tcs) >= 3:
            set_cell_in_tc(tcs[2], value)


def fill_titre_pro(doc):
    """Table 1 — titre professionnel visé + modalité (Parcours coché)."""
    t = doc.tables[1]
    rows = t._tbl.tr_lst
    # row 0 = Titre professionnel visé (libellé)
    # row 1 ou 2 = champ de saisie pour le titre
    # row 6 = ☐ Parcours de formation
    # row 7 = ☐ VAE
    # On écrit le titre dans la première row libre après row 0
    # Row 2 contient le placeholder "Cliquez ici pour entrer l'intitulé…"
    # On l'écrase directement (et au passage, row 1 reste vide pour
    # l'aération du document)
    if len(rows) > 2:
        tcs = rows[2].xpath('.//w:tc')
        if tcs:
            set_cell_in_tc(tcs[0], TITRE_PRO, bold=True)

    # Cocher la case « Parcours de formation » :
    # row 6 = ☐ | Parcours de formation
    # row 7 = ☐ | VAE
    # On remplace le ☐ par ☒ dans la cellule 0 de row 6.
    if len(rows) > 6:
        tcs = rows[6].xpath('.//w:tc')
        if tcs:
            set_cell_in_tc(tcs[0], '☒', bold=True)


def find_table_by_text(doc, needle, skip_indices=()):
    """Trouve la première table dont une cellule contient `needle`,
    en excluant les indices passés dans `skip_indices` (utile pour
    éviter de tomber sur la table Sommaire qui contient un index
    textuel de tous les autres titres).
    """
    for i, t in enumerate(doc.tables):
        if i in skip_indices:
            continue
        for row in t.rows:
            for c in row._tr.xpath('.//w:tc'):
                from lxml import etree
                text = ''.join(etree.tostring(t_el, method='text',
                                              encoding='unicode')
                               for t_el in c.xpath('.//w:t'))
                if needle in text:
                    return t
    return None


def fill_diplomes(doc):
    """Table « Titres, diplômes, CQP » — trouvée par contenu en excluant
    la table Sommaire (qui contient une référence textuelle à ce titre)."""
    # On exclut la Sommaire (tables[3]) qui contient le texte « Titres,
    # diplômes » en tant qu'entrée de table des matières.
    t = find_table_by_text(doc, "Titres, diplômes", skip_indices=(3,))
    if t is None:
        return
    rows = t._tbl.tr_lst
    # row 3 = ligne d'entête (Intitulé | Autorité | Date)
    # rows 4+ = lignes à remplir
    for i, (intitule, autorite, date) in enumerate(DIPLOMES):
        row_idx = 4 + i
        if row_idx >= len(rows):
            break
        tcs = rows[row_idx].xpath('.//w:tc')
        if len(tcs) >= 3:
            set_cell_in_tc(tcs[0], intitule)
            set_cell_in_tc(tcs[1], autorite)
            set_cell_in_tc(tcs[2], date)


def fill_declaration(doc):
    """Paragraphes de la déclaration sur l'honneur — laissés vides
    intentionnellement.

    Le candidat imprime le DP et REMPLIT À LA MAIN sa déclaration
    (nom, ville, date). On remplace donc tous les placeholders du
    modèle (« [prénom et nom] », « [date à compléter] »,
    « Cliquez ici pour taper du texte », « Cliquez ici pour choisir
    une date ») par des lignes de pointillés ou par du vide.
    """
    from docx.oxml.ns import qn
    LINE = "................................................"
    SHORT_LINE = ".................."

    def collapse_runs_text(paragraph):
        """Renvoie le texte complet d'un paragraphe (tous w:t concaténés)."""
        return ''.join(t.text or '' for t in paragraph._p.iter(qn('w:t')))

    def set_paragraph_text(paragraph, new_text):
        """Remplace le contenu textuel du paragraphe : vide tous les
        w:t existants sauf le premier, qui reçoit `new_text`."""
        ts = list(paragraph._p.iter(qn('w:t')))
        if not ts:
            return
        ts[0].text = new_text
        for t in ts[1:]:
            t.text = ''

    for p in doc.paragraphs:
        full = collapse_runs_text(p)
        # Paragraphes contenant uniquement des placeholders SDT
        # (« Cliquez ici pour taper du texte. » et/ou « Cliquez ici pour
        # choisir une date » répétés) → on vide tous les w:t.
        cleaned = full
        for ph in ("Cliquez ici pour taper du texte.",
                   "Cliquez ici pour choisir une date"):
            cleaned = cleaned.replace(ph, "")
        if not cleaned.strip() and 'Cliquez ici' in full:
            set_paragraph_text(p, "")
            continue
        # « Je soussigné(e) [prénom et nom] , » → ligne pointillée
        # On retire la virgule trainante du modèle qui produisait une
        # « ,........... » disgracieuse après le nom.
        if 'Je soussigné' in full and '[prénom et nom]' in full:
            new = full.replace('[prénom et nom]', f' {LINE} ')
            # Supprime la virgule isolée en fin de ligne si présente.
            new = new.rstrip().rstrip(',').rstrip() + ' '
            set_paragraph_text(p, new)
            continue
        # « Fait à … le … » : la ligne contient déjà nos pointillés
        # depuis l'édition précédente. Si on retrouve les anciens
        # placeholders du modèle on les nettoie aussi.
        if full.startswith('Fait à'):
            new = full
            new = new.replace('[date à compléter par le candidat]', SHORT_LINE)
            new = new.replace('Cliquez ici pour choisir une date', SHORT_LINE)
            new = new.replace('Cliquez ici pour taper du texte.', '')
            # S'il n'y a pas encore de pointillés (premier passage sur le
            # modèle d'origine), on injecte une forme propre.
            if '......' not in new:
                new = f'Fait à {SHORT_LINE} le {SHORT_LINE}'
            set_paragraph_text(p, new)


def ensure_page_break_before(doc, needle):
    """Insère un saut de page AVANT le premier paragraphe contenant
    `needle`, sauf si un PAGE_BREAK est déjà présent juste avant.

    Utile pour s'assurer qu'un titre de section (ex : « Exemples de
    pratique professionnelle ») démarre en haut d'une nouvelle page
    plutôt que de se faire couper en bas de page précédente.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for p in doc.paragraphs:
        if needle not in p.text:
            continue
        # Vérifier si un PAGE_BREAK existe déjà dans le paragraphe précédent
        prev = p._p.getprevious()
        if prev is not None:
            brs = prev.xpath('.//w:br[@w:type="page"]')
            if brs:
                return  # déjà présent
        # Insérer un saut de page en début de ce paragraphe
        # via un run avec <w:br w:type="page"/>
        new_run = OxmlElement('w:r')
        new_br = OxmlElement('w:br')
        new_br.set(qn('w:type'), 'page')
        new_run.append(new_br)
        # On insère le run en TÊTE du paragraphe (avant les autres runs)
        p._p.insert(0, new_run)
        return


def remove_excess_page_breaks(doc):
    """Supprime les paragraphes-saut-de-page entre AT1↔AT2 et AT2↔Titres.

    Le modèle force un PAGE_BREAK entre chaque AT pour qu'ils commencent
    sur une nouvelle page. Mais avec du contenu dense (notre cas), cela
    crée des pages quasi vides. On laisse Word/LibreOffice gérer le flux
    naturel.
    """
    from lxml import etree
    body = doc.element.body
    children = list(body)

    # On veut supprimer les paragraphes-saut-de-page qui se trouvent
    # juste avant des tables AT2, Titres-diplômes, Déclaration, etc.
    # Repérage par contenu de la table suivante.
    targets_following = (
        'Activité-type',          # AT2 (AT3 déjà supprimé)
        'Titres, diplômes',       # Table 7 (devenue 6 après suppression AT3)
        # Pas pour Déclaration / Documents / Annexes : on garde leur séparation
    )
    seen_first_at = False

    for i, child in enumerate(children):
        tag = etree.QName(child).localname
        if tag != 'tbl':
            continue
        first_text = ''.join(child.xpath('.//w:t/text()'))[:60]
        if not any(t in first_text for t in targets_following):
            continue
        if 'Activité-type' in first_text and not seen_first_at:
            # première Activité-type rencontrée = AT1, on ne touche pas
            seen_first_at = True
            continue
        # Chercher le paragraphe précédent qui contient un PAGE_BREAK
        prev = child.getprevious()
        while prev is not None:
            if etree.QName(prev).localname == 'p' and prev.xpath('.//w:br[@w:type="page"]'):
                prev.getparent().remove(prev)
                break
            prev = prev.getprevious()


def collapse_double_page_breaks(doc):
    """Supprime le saut de page orphelin entre AT2 ex3 et Titres-diplômes.

    Le template AFPA contenait des sauts de page liés aux activités-type
    AT3/AT4 désormais supprimées. Il en reste un juste avant « Titres,
    diplômes » qui crée une page p.20 entièrement vide. On le retire
    chirurgicalement (sans toucher aux sauts qui séparent AT1↔AT2, ni
    aux sauts insérés par clone_example_table pour les exemples).
    """
    from lxml import etree
    body = doc.element.body
    children = list(body)

    def is_p_with_pgbr(el):
        return (etree.QName(el).localname == 'p'
                and bool(el.xpath('.//w:br[@w:type="page"]')))

    # Repérer la table « Titres, diplômes »
    titres_idx = None
    for i, child in enumerate(children):
        if etree.QName(child).localname != 'tbl':
            continue
        txt = ''.join(child.xpath('.//w:t/text()'))[:60]
        if 'Titres, diplômes' in txt:
            titres_idx = i
            break
    if titres_idx is None:
        return

    # Remonter jusqu'à la table précédente et nettoyer le gap.
    prev_tbl_idx = None
    for j in range(titres_idx - 1, -1, -1):
        if etree.QName(children[j]).localname == 'tbl':
            prev_tbl_idx = j
            break
    if prev_tbl_idx is None:
        return

    gap_indices = list(range(prev_tbl_idx + 1, titres_idx))
    pgbr_in_gap = [k for k in gap_indices if is_p_with_pgbr(children[k])]
    for k in pgbr_in_gap:
        p = children[k]
        if p.getparent() is not None:
            p.getparent().remove(p)


def remove_at3_section(doc):
    """Supprime la table AT3 (Activité-type 3) + son saut de page précédent.

    Recherche par contenu : la table dont la première cellule contient
    « Activité-type » avec le numéro 3 ou son intitulé non-applicable.
    """
    body = doc.element.body

    # Identifier la table AT3 (et son saut de page précédent)
    target_table = None
    for child in body:
        from lxml import etree
        if etree.QName(child).localname != 'tbl':
            continue
        first_text = ''.join(child.xpath('.//w:t/text()'))[:60]
        # AT3 contient « Activité-type » et le numéro 3 (du sdt comboBox)
        # ou notre marqueur « Non applicable »
        if 'Activité-type' in first_text and ('3' in first_text or 'Non applicable' in first_text):
            # Vérifier que c'est bien AT3 (pas AT1 ou AT2)
            if 'front-end' not in first_text and 'back-end' not in first_text:
                target_table = child
                break

    if target_table is None:
        return  # rien à faire

    # Trouver le paragraphe juste avant la table (qui contient le PAGE_BREAK)
    prev_para = target_table.getprevious()
    while prev_para is not None and etree.QName(prev_para).localname != 'p':
        prev_para = prev_para.getprevious()

    # Supprimer la table
    target_table.getparent().remove(target_table)

    # Supprimer le saut de page précédent s'il en contient un
    if prev_para is not None and prev_para.xpath('.//w:br[@w:type="page"]'):
        prev_para.getparent().remove(prev_para)


def fill_sommaire(doc):
    """Table 3 — Sommaire.

    On remplit les intitulés AT1/AT2 + exemples. AT3 et AT4 marqués
    « Non applicable » (DWWM = 2 AT seulement). Les numéros de page
    réels seront déterminés au moment de l'impression — on laisse
    des indicateurs approximatifs.
    """
    t = doc.tables[3]
    rows = t._tbl.tr_lst

    def write_cell(row_idx, col_idx, text):
        if row_idx >= len(rows):
            return
        tcs = rows[row_idx].xpath('.//w:tc')
        if not tcs:
            return
        # Support négatif (col_idx = -1 → dernière colonne).
        if col_idx < 0:
            col_idx = len(tcs) + col_idx
        if col_idx < 0 or col_idx >= len(tcs):
            return
        set_cell_in_tc(tcs[col_idx], text)

    # Numéros mesurés sur le PDF final (29 pages, layout v2026-06) :
    #   p 1  Identité
    #   p 2  Présentation du dossier
    #   p 3  Sommaire
    #   p 4  Titre « EXEMPLES DE PRATIQUE PROFESSIONNELLE »
    #   p 5  AT1 ex 1 (front Crew React 18 + AuthContext)
    #   p 8  AT1 ex 2 (grille Planning + drag-and-drop)
    #   p 11 AT1 ex 3 (Smart Planner modale)
    #   p 13 AT2 ex 1 (Solver HCR Node.js/Express)
    #   p 16 AT2 ex 2 (Auth JWT + middlewares)
    #   p 18 AT2 ex 3 (iCalendar RFC 5545)
    #   p 20 Titres, diplômes
    #   p 21 Déclaration sur l'honneur
    #   p 22 Documents illustrant la pratique
    #   p 23 Annexes (titre) + p24-29 captures d'écran

    # AT1 (rows 2-5)
    write_cell(2, 0, AT1_INTITULE);          write_cell(2, -1, "5")
    write_cell(3, 1, AT1_EX_INTITULE);       write_cell(3, -1, "5")
    write_cell(4, 1, AT1B_EX_INTITULE);      write_cell(4, -1, "8")
    write_cell(5, 1, AT1C_EX_INTITULE);      write_cell(5, -1, "11")

    # AT2 (rows 7-10)
    write_cell(7, 0, AT2_INTITULE);          write_cell(7, -1, "13")
    write_cell(8, 1, AT2_EX_INTITULE);       write_cell(8, -1, "13")
    write_cell(9, 1, AT2B_EX_INTITULE);      write_cell(9, -1, "16")
    write_cell(10, 1, AT2C_EX_INTITULE);     write_cell(10, -1, "18")

    # AT3 et AT4 sont non applicables au titre DWWM. On supprime
    # entièrement les 10 rangées correspondantes (11-20) du sommaire
    # pour que le jury ne voie plus de placeholders vides.
    # Suppression de la fin vers le début pour préserver les indices.
    for idx in range(20, 10, -1):
        if idx < len(rows):
            tr = rows[idx]
            tr.getparent().remove(tr)
    # Refresh rows ref pour les écritures suivantes.
    rows = t._tbl.tr_lst

    # Diplômes / Déclaration / Documents / Annexes — après suppression
    # des 10 rangées AT3/AT4 (11-20), ces rangées sont désormais en
    # indices 12, 13, 14, 15 (au lieu de 22-25).
    write_cell(12, -1, "20")  # Titres, diplômes
    write_cell(13, -1, "21")  # Déclaration sur l'honneur
    write_cell(14, -1, "22")  # Documents illustrant
    write_cell(15, 0, "Annexes — 6 captures d'écran représentatives "
                      "de l'application Crew (pages 24-29)")
    write_cell(15, -1, "23")


def fill_documents_illustrant(doc):
    """Table « Documents illustrant la pratique professionnelle » —
    trouvée par contenu en excluant la Sommaire qui contient une
    référence textuelle à ce titre.
    """
    t = find_table_by_text(doc, "Documents illustrant la pratique",
                           skip_indices=(3,))
    if t is None:
        return
    rows = t._tbl.tr_lst
    # row 3 = "Intitulé" + placeholder "Cliquez ici…"
    # rows 4+ = à remplir
    documents = [
        "Code source du projet Crew — "
            "github.com/Vladimir08880888/crew",
        "Application déployée en production — "
            "crew-planner-hazel.vercel.app",
        "Repo centralisé des dossiers AFPA — "
            "github.com/Vladimir08880888/vladimir-rodzaevski-dossiers-pros-DWWM",
        "Dossier Professionnel détaillé (3 exemples par AT, code commenté) — "
            "dossier-professionnel/DP.pdf",
        "Cahier des charges (sections 1-11, évolutions v2 documentées) — "
            "dossier-projet/cahier-des-charges.md",
        "Justification scientifique (12 références peer-reviewed — INRS, "
            "NIOSH, Convention HCR, OR economics) — "
            "dossier-projet/annexes/JUSTIFICATION_SCIENTIFIQUE.md",
        "Schéma BDD (Mermaid ERD + règles d'intégrité) — "
            "dossier-projet/annexes/SCHEMA_BDD.md",
        "Manuel utilisateur (manager + équipier) — "
            "dossier-projet/annexes/MANUEL_UTILISATEUR.md",
        "17 captures d'écran haute résolution — "
            "dossier-projet/annexes/screenshots/",
        "Plan de tests (60+ scénarios Playwright, 29/29 PASS sur v2) — "
            "dossier-projet/annexes/PLAN_DE_TESTS.md",
        "Trame de soutenance — "
            "dossier-projet/soutenance.md",
    ]
    for i, doc_str in enumerate(documents):
        row_idx = 3 + i
        if row_idx >= len(rows):
            break
        tcs = rows[row_idx].xpath('.//w:tc')
        if not tcs:
            continue
        set_cell_in_tc(tcs[0], doc_str)

# ────────────────────────────────────────────────────────────────────────
# Captures d'écran à insérer dans la section « Annexes » du DP
# ────────────────────────────────────────────────────────────────────────
SCREENSHOTS = [
    ("04-dashboard-manager.png", 5.8,
     "Tableau de bord du manager — heures planifiées par équipier vs cibles, "
     "couverture par poste, alertes santé du service."),
    ("07-planning-grid.png", 5.8,
     "Grille de planning hebdomadaire — équipiers groupés par poste "
     "(cuisine, salle), ligne « Extras » avec déficits à couvrir, "
     "indicateurs de couverture par service en en-tête, masse salariale "
     "prévisionnelle affichée."),
    ("08-smart-planner-modal.png", 5.8,
     "Modale Smart Planner — tableau (jour × service) de densité prévue "
     "avec presets Calme 0,5 / Normal 1,0 / Chargé 1,3, couverture par "
     "service et par poste, liste des services non couverts."),
    ("06-member-edit-modal.png", 5.8,
     "Modale d'édition équipier — 5 profils nommés (Apprenti → Référent), "
     "polyvalence multi-postes, heures cibles, permissions."),
    ("09-stats-charts.png", 5.8,
     "Page Statistiques — graphiques Chart.js (heures planifiées vs cibles, "
     "couverture par poste, évolution hebdomadaire)."),
    # Capture mobile : largeur très réduite (1,5 inch) car le smartphone
    # est en format portrait étroit (390×2606 px). Sans réduction forte,
    # l'image déborde du bas de page et la légende part sur une page
    # blanche orpheline.
    ("15-mobile-dashboard.png", 1.5,
     "Vue mobile responsive — application utilisable sur smartphone."),
]


def clone_example_table(doc, source_table_idx):
    """Clone une table d'exemple (Table 4 ou 5) et l'insère juste après
    la source, précédée d'un saut de page. Retourne le nouvel objet Table
    python-docx.

    Stratégie : deepcopy de l'élément <w:tbl> source puis insertion via
    le parent body element. Tous les indices doc.tables[N] postérieurs
    glissent de +1.

    Le saut de page inséré juste avant la clone empêche LibreOffice de
    coincer les premières lignes (en-tête « Activité-type X / Exemple
    n°Y / intitulé / Compétences couvertes ») en bas de la page
    précédente, en laissant flotter le corps en haut de la suivante.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    source = doc.tables[source_table_idx]
    src_el = source._tbl
    new_el = deepcopy(src_el)
    src_el.addnext(new_el)
    # Saut de page entre la source et la clone : <w:p><w:r><w:br type=page/></w:r></w:p>
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    new_el.addprevious(pb_p)
    return doc.tables[source_table_idx + 1]


def append_visual_annexes(doc):
    """Ajoute les captures d'écran représentatives sous la section « Annexes ».

    La table 9 (« Annexes ») existe en fin de document avec un en-tête
    mais aucun contenu. On ajoute les screenshots juste APRÈS cette table,
    en pleine largeur (6 inches ≈ 15 cm) avec une légende italique.
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    screens_dir = (DST.parent.parent / "dossier-projet" / "annexes"
                   / "screenshots")
    if not screens_dir.exists():
        # Fallback : essayer le repo crew si on est dans son arborescence.
        return

    # Petit titre de section avant les images.
    h = doc.add_paragraph()
    r = h.add_run("Captures d'écran représentatives de l'application Crew")
    r.bold = True
    r.font.size = Pt(13)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for entry in SCREENSHOTS:
        if len(entry) == 3:
            fname, width_in, caption = entry
        else:
            fname, caption = entry
            width_in = 5.8
        img = screens_dir / fname
        if not img.exists():
            continue
        # Image + légende dans un même paragraphe (séparées par un saut
        # de ligne intra-paragraphe), pour que LibreOffice ne sépare
        # jamais la légende de son image sur deux pages.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Garder le paragraphe entier groupé avec le suivant.
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(img), width=Inches(width_in))
        # Saut de ligne intra-paragraphe puis légende italique.
        run.add_break()
        cap_run = p.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)


# ────────────────────────────────────────────────────────────────────────
# MAIN
# ────────────────────────────────────────────────────────────────────────

def main():
    if not SRC.exists():
        raise SystemExit(f"Source introuvable : {SRC}")
    DST.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(SRC))

    fill_header_table(doc)
    fill_titre_pro(doc)

    # AT1 → Table 4 (exemple 1) + 2 clones (exemples 2 et 3).
    # On clone AVANT de remplir pour conserver l'élément vierge.
    at1_ex2 = clone_example_table(doc, 4)
    at1_ex3 = clone_example_table(doc, 5)  # 5 = position après ex2

    fill_example_table(
        doc.tables[4],
        intitule_at=AT1_INTITULE,
        intitule_ex=AT1_EX_INTITULE,
        tasks=AT1_TASKS,
        moyens=AT1_MOYENS,
        avec_qui=AT1_AVEC_QUI,
        nom_ent=AT1_NOM_ENTREPRISE,
        service=AT1_SERVICE,
        periode_du=AT1_PERIODE_DU,
        periode_au=AT1_PERIODE_AU,
        infos_compl=AT1_INFOS_COMPL, exemple_num=1,
    )
    fill_example_table(
        at1_ex2,
        intitule_at=AT1_INTITULE,
        intitule_ex=AT1B_EX_INTITULE,
        tasks=AT1B_TASKS,
        moyens=AT1B_MOYENS,
        avec_qui=AT1B_AVEC_QUI,
        nom_ent=AT1_NOM_ENTREPRISE,
        service=AT1_SERVICE,
        periode_du=AT1_PERIODE_DU,
        periode_au=AT1_PERIODE_AU,
        infos_compl="", exemple_num=2,
    )
    fill_example_table(
        at1_ex3,
        intitule_at=AT1_INTITULE,
        intitule_ex=AT1C_EX_INTITULE,
        tasks=AT1C_TASKS,
        moyens=AT1C_MOYENS,
        avec_qui=AT1C_AVEC_QUI,
        nom_ent=AT1_NOM_ENTREPRISE,
        service=AT1_SERVICE,
        periode_du=AT1_PERIODE_DU,
        periode_au=AT1_PERIODE_AU,
        infos_compl="", exemple_num=3,
    )

    # AT2 → Table 7 (après les 2 clones AT1) + 2 clones.
    # Index courant AT2 = 4 + 3 (originale + 2 clones) = 7.
    at2_idx = 7
    at2_ex2 = clone_example_table(doc, at2_idx)
    at2_ex3 = clone_example_table(doc, at2_idx + 1)

    fill_example_table(
        doc.tables[at2_idx],
        intitule_at=AT2_INTITULE,
        intitule_ex=AT2_EX_INTITULE,
        tasks=AT2_TASKS,
        moyens=AT2_MOYENS,
        avec_qui=AT2_AVEC_QUI,
        nom_ent=AT2_NOM_ENTREPRISE,
        service=AT2_SERVICE,
        periode_du=AT2_PERIODE_DU,
        periode_au=AT2_PERIODE_AU,
        infos_compl=AT2_INFOS_COMPL, exemple_num=1,
    )
    fill_example_table(
        at2_ex2,
        intitule_at=AT2_INTITULE,
        intitule_ex=AT2B_EX_INTITULE,
        tasks=AT2B_TASKS,
        moyens=AT2B_MOYENS,
        avec_qui=AT2B_AVEC_QUI,
        nom_ent=AT2_NOM_ENTREPRISE,
        service=AT2_SERVICE,
        periode_du=AT2_PERIODE_DU,
        periode_au=AT2_PERIODE_AU,
        infos_compl="", exemple_num=2,
    )
    fill_example_table(
        at2_ex3,
        intitule_at=AT2_INTITULE,
        intitule_ex=AT2C_EX_INTITULE,
        tasks=AT2C_TASKS,
        moyens=AT2C_MOYENS,
        avec_qui=AT2C_AVEC_QUI,
        nom_ent=AT2_NOM_ENTREPRISE,
        service=AT2_SERVICE,
        periode_du=AT2_PERIODE_DU,
        periode_au=AT2_PERIODE_AU,
        infos_compl="", exemple_num=3,
    )

    # AT3 → laissée en place temporairement pour stabilité d'indexation
    # (sera supprimée à la toute fin).

    fill_sommaire(doc)
    fill_diplomes(doc)
    fill_declaration(doc)
    fill_documents_illustrant(doc)

    # En dernier — supprimer la table AT3 (DWWM = 2 AT seulement).
    # Doit être après tous les fills pour que les indices doc.tables[N]
    # restent stables.
    remove_at3_section(doc)

    # NOTE : on n'appelle PAS remove_excess_page_breaks() — la suppression
    # systématique des sauts de page entre AT1↔AT2 et entre les exemples
    # provoque des en-têtes coupés en bas de page (le titre « Activité-type 2 »
    # se retrouvait par ex. en bas de la page 10 au lieu d'ouvrir la 11).
    # On laisse chaque table d'exemple démarrer en haut de sa propre page.
    #
    # Le template AFPA contient déjà un saut de section avant le titre
    # « EXEMPLES DE PRATIQUE PROFESSIONNELLE » — pas besoin d'en forcer
    # un supplémentaire (sinon page blanche p.4).
    #
    # En revanche, le template comportait des doubles sauts de page
    # destinés à séparer AT3/AT4 (qu'on a supprimés) — d'où une p.20
    # vide juste avant « Titres, diplômes ». On consolide.
    collapse_double_page_breaks(doc)

    append_visual_annexes(doc)

    doc.save(str(DST))
    print(f"✓ Généré : {DST}")
    print(f"  Taille  : {DST.stat().st_size:,} octets")


if __name__ == '__main__':
    main()
